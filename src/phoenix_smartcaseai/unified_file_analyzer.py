"""
Unified File Analysis Module for Phoenix-SmartCaseAI

This module provides comprehensive file analysis with intelligent LLM routing.
Combines foundational file processing with AI-powered analysis and smart provider routing.
"""

import os
import json
import xml.etree.ElementTree as ET
import base64
import mimetypes
from pathlib import Path
from typing import Dict, List, Optional, Union, Any

# Document processing
try:
    import docx
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

# Image processing
try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# LLM providers
try:
    from langchain_openai import ChatOpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

try:
    from langchain_google_genai import ChatGoogleGenerativeAI
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False

try:
    from langchain_anthropic import ChatAnthropic
    CLAUDE_AVAILABLE = True
except ImportError:
    CLAUDE_AVAILABLE = False


class FileAnalysisResult:
    """Container for file analysis results."""
    
    def __init__(self, file_path: str, file_type: str, content: str, metadata: Dict[str, Any]):
        self.file_path = file_path
        self.file_type = file_type
        self.content = content
        self.metadata = metadata
        self.summary = self._generate_summary()
    
    def _generate_summary(self) -> str:
        """Generate a summary of the file content."""
        if not self.content:
            return "No content extracted from file."
        
        # Truncate content for summary
        content_preview = self.content[:500] + "..." if len(self.content) > 500 else self.content
        return f"File type: {self.file_type}\nContent preview: {content_preview}"
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for serialization."""
        return {
            "file_path": self.file_path,
            "file_type": self.file_type,
            "content": self.content,
            "metadata": self.metadata,
            "summary": self.summary
        }


class UnifiedFileAnalyzer:
    """
    Unified file analyzer with intelligent LLM routing and comprehensive file processing.
    
    Features:
    - Supports all major file types (text, PDF, Word, Excel, images, videos, JSON, XML)
    - Intelligent routing: Images/videos → Gemini, Documents → OpenAI/Claude
    - OCR for images, structured data extraction
    - LLM-enhanced analysis for deeper insights
    """
    
    def __init__(self, 
                 openai_api_key: Optional[str] = None,
                 google_api_key: Optional[str] = None, 
                 anthropic_api_key: Optional[str] = None):
        """
        Initialize the unified file analyzer.
        
        Args:
            openai_api_key: OpenAI API key for text analysis and vision
            google_api_key: Google API key for Gemini (best for images/videos)
            anthropic_api_key: Anthropic API key for Claude analysis
        """
        self.openai_api_key = openai_api_key or os.getenv("OPENAI_API_KEY")
        self.google_api_key = google_api_key or os.getenv("GOOGLE_API_KEY")
        self.anthropic_api_key = anthropic_api_key or os.getenv("ANTHROPIC_API_KEY")
        
        # File type classifications
        self.visual_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp', '.tiff'}
        self.document_extensions = {'.txt', '.md', '.pdf', '.docx', '.doc', '.rtf'}
        self.data_extensions = {'.json', '.xml', '.csv', '.xlsx', '.xls'}
        self.video_extensions = {'.mp4', '.avi', '.mov', '.wmv', '.mkv', '.webm'}
        
        # Initialize LLM clients
        self.llm_clients = {}
        self._initialize_llm_clients()
    
    def _initialize_llm_clients(self):
        """Initialize available LLM clients based on API keys."""
        if self.openai_api_key and OPENAI_AVAILABLE:
            try:
                self.llm_clients["openai"] = ChatOpenAI(
                    model="gpt-4o-mini",
                    api_key=self.openai_api_key,
                    temperature=0.3
                )
            except Exception as e:
                print(f"Warning: Failed to initialize OpenAI: {e}")
        
        if self.google_api_key and GEMINI_AVAILABLE:
            try:
                self.llm_clients["gemini"] = ChatGoogleGenerativeAI(
                    model="gemini-2.5-flash",
                    google_api_key=self.google_api_key,
                    temperature=0.3
                )
            except Exception as e:
                print(f"Warning: Failed to initialize Gemini: {e}")
        
        if self.anthropic_api_key and CLAUDE_AVAILABLE:
            try:
                self.llm_clients["claude"] = ChatAnthropic(
                    model="claude-3-5-haiku-20241022",
                    api_key=self.anthropic_api_key,
                    temperature=0.3
                )
            except Exception as e:
                print(f"Warning: Failed to initialize Claude: {e}")
    
    def _get_provider_for_file(self, file_path: Union[str, Path]) -> str:
        """Determine the best LLM provider for a given file type."""
        file_ext = Path(file_path).suffix.lower()
        
        # Images and videos: prefer Gemini for visual analysis
        if file_ext in self.visual_extensions or file_ext in self.video_extensions:
            if "gemini" in self.llm_clients:
                return "gemini"
            elif "openai" in self.llm_clients:
                return "openai"  # OpenAI has vision capabilities too
        
        # Documents and data: prefer OpenAI or Claude
        if "openai" in self.llm_clients:
            return "openai"
        elif "claude" in self.llm_clients:
            return "claude"
        elif "gemini" in self.llm_clients:
            return "gemini"
        
        return "none"
    
    def analyze_file(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Analyze a single file with intelligent LLM routing.
        
        Args:
            file_path: Path to the file to analyze
            
        Returns:
            Dictionary containing analysis results
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {
                "file_path": str(file_path),
                "file_type": "error",
                "error": "File not found",
                "content": "",
                "analysis": "File not found"
            }
        
        # Get basic file analysis
        try:
            base_result = self._extract_content_and_metadata(file_path)
        except Exception as e:
            return {
                "file_path": str(file_path),
                "file_type": "error", 
                "error": str(e),
                "content": "",
                "analysis": f"Failed to extract content: {e}"
            }
        
        # Route to appropriate LLM for enhanced analysis
        provider = self._get_provider_for_file(file_path)
        
        if provider != "none":
            try:
                enhanced_analysis = self._enhance_with_llm(base_result, provider)
            except Exception as e:
                print(f"Warning: LLM enhancement failed for {file_path}: {e}")
                enhanced_analysis = base_result.summary
        else:
            enhanced_analysis = base_result.summary
        
        return {
            "file_path": str(file_path),
            "file_type": base_result.file_type,
            "content": base_result.content,
            "metadata": base_result.metadata,
            "analysis": enhanced_analysis,
            "provider": provider,
            "enhanced_analysis": enhanced_analysis
        }
    
    def _extract_content_and_metadata(self, file_path: Path) -> FileAnalysisResult:
        """Extract content and metadata from file based on its type."""
        mime_type, _ = mimetypes.guess_type(str(file_path))
        file_ext = file_path.suffix.lower()
        
        # Route to appropriate extraction method
        if file_ext in {'.txt', '.md', '.py', '.js', '.html', '.css'}:
            return self._analyze_text_file(file_path, mime_type or 'text/plain')
        elif file_ext == '.pdf':
            return self._analyze_pdf_file(file_path, mime_type or 'application/pdf')
        elif file_ext in {'.docx', '.doc'}:
            return self._analyze_docx_file(file_path, mime_type or 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        elif file_ext in {'.xlsx', '.xls', '.csv'}:
            return self._analyze_spreadsheet_file(file_path, mime_type or 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        elif file_ext == '.json':
            return self._analyze_json_file(file_path, mime_type or 'application/json')
        elif file_ext == '.xml':
            return self._analyze_xml_file(file_path, mime_type or 'application/xml')
        elif file_ext in self.visual_extensions:
            return self._analyze_image_file(file_path, mime_type or 'image/jpeg')
        elif file_ext in self.video_extensions:
            return self._analyze_video_file(file_path, mime_type or 'video/mp4')
        else:
            return self._analyze_unknown_file(file_path, mime_type or 'application/octet-stream')
    
    def _analyze_text_file(self, file_path: Path, mime_type: str) -> FileAnalysisResult:
        """Analyze text-based files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
            except Exception as e:
                content = f"Error reading file: {str(e)}"
        
        metadata = {
            "file_size": file_path.stat().st_size,
            "lines": len(content.split('\n')) if content else 0,
            "characters": len(content),
            "encoding": "utf-8"
        }
        
        return FileAnalysisResult(
            file_path=str(file_path),
            file_type="text",
            content=content,
            metadata=metadata
        )
    
    def _analyze_pdf_file(self, file_path: Path, mime_type: str) -> FileAnalysisResult:
        """Analyze PDF files using available libraries."""
        content = ""
        metadata = {"file_size": file_path.stat().st_size}
        
        # Try pdfplumber first (better text extraction)
        if PDFPLUMBER_AVAILABLE:
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    metadata["pages"] = len(pdf.pages)
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            content += page_text + "\n"
            except Exception as e:
                content = f"pdfplumber extraction failed: {str(e)}"
        
        # Fallback to PyPDF2
        elif PDF_AVAILABLE:
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    metadata["pages"] = len(reader.pages)
                    for page in reader.pages:
                        content += page.extract_text() + "\n"
            except Exception as e:
                content = f"PyPDF2 extraction failed: {str(e)}"
        else:
            content = "PDF processing libraries not available. Install pdfplumber or PyPDF2."
        
        return FileAnalysisResult(
            file_path=str(file_path),
            file_type="pdf",
            content=content.strip(),
            metadata=metadata
        )
    
    def _analyze_docx_file(self, file_path: Path, mime_type: str) -> FileAnalysisResult:
        """Analyze Word documents."""
        content = ""
        metadata = {"file_size": file_path.stat().st_size}
        
        if DOCX_AVAILABLE:
            try:
                doc = Document(file_path)
                paragraphs = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text)
                
                content = "\n".join(paragraphs)
                metadata["paragraphs"] = len(paragraphs)
                metadata["tables"] = len(doc.tables)
            except Exception as e:
                content = f"DOCX extraction failed: {str(e)}"
        else:
            content = "python-docx library not available. Install python-docx."
        
        return FileAnalysisResult(
            file_path=str(file_path),
            file_type="docx",
            content=content,
            metadata=metadata
        )
    
    def _analyze_spreadsheet_file(self, file_path: Path, mime_type: str) -> FileAnalysisResult:
        """Analyze spreadsheet files (Excel, CSV)."""
        content = ""
        metadata = {"file_size": file_path.stat().st_size}
        
        if PANDAS_AVAILABLE:
            try:
                if file_path.suffix.lower() == '.csv':
                    df = pd.read_csv(file_path, nrows=100)  # Limit to first 100 rows
                else:
                    df = pd.read_excel(file_path, nrows=100)  # Limit to first 100 rows
                
                metadata["rows"] = len(df)
                metadata["columns"] = len(df.columns)
                metadata["column_names"] = list(df.columns)
                
                # Create summary of data
                content = f"Spreadsheet Summary:\n"
                content += f"Columns: {', '.join(df.columns)}\n"
                content += f"Shape: {df.shape}\n\n"
                content += "Sample data:\n"
                content += df.head(10).to_string()
                
            except Exception as e:
                content = f"Spreadsheet extraction failed: {str(e)}"
        else:
            content = "pandas library not available. Install pandas."
        
        return FileAnalysisResult(
            file_path=str(file_path),
            file_type="spreadsheet",
            content=content,
            metadata=metadata
        )
    
    def _analyze_json_file(self, file_path: Path, mime_type: str) -> FileAnalysisResult:
        """Analyze JSON files."""
        content = ""
        metadata = {"file_size": file_path.stat().st_size}
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Analyze structure
            structure = self._analyze_json_structure(data)
            metadata.update(structure)
            
            # Create readable content
            content = json.dumps(data, indent=2, ensure_ascii=False)
            if len(content) > 5000:
                content = content[:5000] + "\n... [truncated]"
                
        except Exception as e:
            content = f"JSON parsing failed: {str(e)}"
        
        return FileAnalysisResult(
            file_path=str(file_path),
            file_type="json",
            content=content,
            metadata=metadata
        )
    
    def _analyze_json_structure(self, data: Any, path: str = "") -> Dict[str, Any]:
        """Analyze JSON structure recursively."""
        if isinstance(data, dict):
            return {
                "type": "object",
                "keys": list(data.keys()),
                "key_count": len(data)
            }
        elif isinstance(data, list):
            return {
                "type": "array",
                "length": len(data),
                "item_types": list(set(type(item).__name__ for item in data[:10]))
            }
        else:
            return {"type": type(data).__name__}
    
    def _analyze_xml_file(self, file_path: Path, mime_type: str) -> FileAnalysisResult:
        """Analyze XML files."""
        content = ""
        metadata = {"file_size": file_path.stat().st_size}
        
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            # Analyze structure
            structure = self._analyze_xml_structure(root)
            metadata.update(structure)
            metadata["root_tag"] = root.tag
            
            # Create readable content
            content = ET.tostring(root, encoding='unicode')
            if len(content) > 5000:
                content = content[:5000] + "\n... [truncated]"
                
        except Exception as e:
            content = f"XML parsing failed: {str(e)}"
        
        return FileAnalysisResult(
            file_path=str(file_path),
            file_type="xml",
            content=content,
            metadata=metadata
        )
    
    def _analyze_xml_structure(self, element: ET.Element) -> Dict[str, Any]:
        """Analyze XML structure."""
        children = list(element)
        return {
            "tag": element.tag,
            "attributes": element.attrib,
            "child_count": len(children),
            "child_tags": list(set(child.tag for child in children))
        }
    
    def _analyze_image_file(self, file_path: Path, mime_type: str) -> FileAnalysisResult:
        """Analyze image files with OCR if available."""
        content = ""
        metadata = {"file_size": file_path.stat().st_size}
        
        try:
            if OCR_AVAILABLE:
                image = Image.open(file_path)
                metadata["dimensions"] = image.size
                metadata["mode"] = image.mode
                
                # Try OCR
                try:
                    ocr_text = pytesseract.image_to_string(image)
                    if ocr_text.strip():
                        content = f"OCR extracted text:\n{ocr_text}"
                    else:
                        content = "No text detected in image"
                except Exception as e:
                    content = f"OCR failed: {str(e)}"
            else:
                content = "Image analysis libraries not available. Install Pillow and pytesseract."
        except Exception as e:
            content = f"Image processing failed: {str(e)}"
        
        return FileAnalysisResult(
            file_path=str(file_path),
            file_type="image",
            content=content,
            metadata=metadata
        )
    
    def _analyze_video_file(self, file_path: Path, mime_type: str) -> FileAnalysisResult:
        """Analyze video files (basic metadata)."""
        metadata = {"file_size": file_path.stat().st_size}
        content = f"Video file: {file_path.name}\nSize: {metadata['file_size']} bytes"
        
        return FileAnalysisResult(
            file_path=str(file_path),
            file_type="video",
            content=content,
            metadata=metadata
        )
    
    def _analyze_unknown_file(self, file_path: Path, mime_type: str) -> FileAnalysisResult:
        """Handle unknown file types."""
        metadata = {
            "file_size": file_path.stat().st_size,
            "mime_type": mime_type
        }
        content = f"Unknown file type: {file_path.suffix}\nMIME type: {mime_type}"
        
        return FileAnalysisResult(
            file_path=str(file_path),
            file_type="unknown",
            content=content,
            metadata=metadata
        )
    
    def _enhance_with_llm(self, base_result: FileAnalysisResult, provider: str) -> str:
        """Enhance analysis using LLM providers."""
        if provider not in self.llm_clients:
            return base_result.summary
        
        llm = self.llm_clients[provider]
        
        # Create appropriate prompt based on file type
        if base_result.file_type in ["image", "video"]:
            prompt = self._create_visual_analysis_prompt(base_result)
        else:
            prompt = self._create_document_analysis_prompt(base_result)
        
        try:
            response = llm.invoke(prompt)
            return response.content
        except Exception as e:
            print(f"Warning: LLM analysis failed: {e}")
            return base_result.summary
    
    def _create_visual_analysis_prompt(self, result: FileAnalysisResult) -> str:
        """Create prompt for visual content analysis."""
        return f"""
Analyze this visual content for test case generation:

File: {result.file_path}
Type: {result.file_type}
Content: {result.content}
Metadata: {result.metadata}

Please provide insights for:
1. UI/UX elements visible
2. User interaction points
3. Layout and design patterns
4. Accessibility considerations
5. Test scenarios for this visual content

Focus on actionable insights for comprehensive test case generation.
"""
    
    def _create_document_analysis_prompt(self, result: FileAnalysisResult) -> str:
        """Create prompt for document analysis."""
        return f"""
Analyze this document content for test case generation:

File: {result.file_path}
Type: {result.file_type}
Content: {result.content[:2000]}...
Metadata: {result.metadata}

Please provide insights for:
1. Key requirements and specifications
2. Business rules and constraints
3. Data structures and formats
4. Security and compliance considerations
5. Integration points and dependencies

Focus on extracting testable requirements and scenarios.
"""
    
    def analyze_multiple_files(self, file_paths: List[Union[str, Path]]) -> List[Dict[str, Any]]:
        """
        Analyze multiple files and return their analysis results.
        
        Args:
            file_paths: List of file paths to analyze
            
        Returns:
            List of analysis results for each file
        """
        results = []
        for file_path in file_paths:
            try:
                result = self.analyze_file(file_path)
                results.append(result)
            except Exception as e:
                print(f"Warning: Failed to analyze {file_path}: {e}")
                # Continue with other files
        return results
    
    def generate_combined_analysis(self, results: List[Dict[str, Any]]) -> str:
        """
        Generate a combined analysis summary from multiple file analysis results.
        
        Args:
            results: List of file analysis results
            
        Returns:
            Combined analysis summary as a string
        """
        if not results:
            return "No files were successfully analyzed."
        
        combined_summary = "=== COMPREHENSIVE FILE ANALYSIS ===\n\n"
        
        # Group files by type
        file_types = {}
        for result in results:
            file_type = result.get('file_type', 'unknown')
            if file_type not in file_types:
                file_types[file_type] = []
            file_types[file_type].append(result)
        
        # Generate summary by file type
        for file_type, type_results in file_types.items():
            combined_summary += f"## {file_type.upper()} FILES ({len(type_results)} files)\n\n"
            
            for result in type_results:
                file_name = Path(result.get('file_path', 'unknown')).name
                provider = result.get('provider', 'local')
                enhanced_analysis = result.get('enhanced_analysis', result.get('analysis', 'No analysis available'))
                
                combined_summary += f"### {file_name} (analyzed by {provider})\n"
                combined_summary += f"{enhanced_analysis}\n\n"
        
        # Add overall insights
        combined_summary += "## COMPREHENSIVE INSIGHTS\n\n"
        combined_summary += f"- Total files analyzed: {len(results)}\n"
        combined_summary += f"- File types covered: {', '.join(file_types.keys())}\n"
        
        # Extract key themes
        all_analyses = [result.get('enhanced_analysis', result.get('analysis', '')) for result in results]
        combined_text = ' '.join(all_analyses).lower()
        
        # Look for common themes
        themes = []
        if 'security' in combined_text or 'ssl' in combined_text or 'encrypt' in combined_text:
            themes.append("Security requirements")
        if 'payment' in combined_text or 'transaction' in combined_text:
            themes.append("Payment processing")
        if 'api' in combined_text or 'endpoint' in combined_text:
            themes.append("API integration")
        if 'user' in combined_text or 'interface' in combined_text or 'ui' in combined_text:
            themes.append("User interface")
        if 'test' in combined_text or 'validation' in combined_text:
            themes.append("Testing scenarios")
        if 'wireframe' in combined_text or 'mockup' in combined_text:
            themes.append("UI/UX design")
        
        if themes:
            combined_summary += f"- Key themes identified: {', '.join(themes)}\n"
        
        # Provider usage summary
        providers_used = [result.get('provider', 'local') for result in results]
        provider_counts = {p: providers_used.count(p) for p in set(providers_used)}
        combined_summary += f"- LLM providers utilized: {dict(provider_counts)}\n"
        
        combined_summary += "\nThis unified analysis combines foundational file processing with AI-enhanced insights for comprehensive test case generation.\n"
        
        return combined_summary
    
    def get_status_report(self) -> Dict[str, Any]:
        """Get status report of available providers and capabilities."""
        return {
            "available_providers": list(self.llm_clients.keys()),
            "visual_analysis_provider": "gemini" if "gemini" in self.llm_clients else ("openai" if "openai" in self.llm_clients else "none"),
            "document_analysis_providers": [p for p in self.llm_clients.keys() if p in ["openai", "claude"]],
            "supported_visual_formats": list(self.visual_extensions),
            "supported_document_formats": list(self.document_extensions),
            "supported_data_formats": list(self.data_extensions),
            "supported_video_formats": list(self.video_extensions),
            "total_supported_formats": len(self.visual_extensions) + len(self.document_extensions) + len(self.data_extensions) + len(self.video_extensions),
            "ocr_available": OCR_AVAILABLE,
            "pdf_processing": PDFPLUMBER_AVAILABLE or PDF_AVAILABLE,
            "office_docs": DOCX_AVAILABLE,
            "spreadsheets": PANDAS_AVAILABLE
        }
