"""
File Analysis Module for Phoenix-SmartCaseAI

This module provides comprehensive file analysis capabilities for various file types
including documents, images, videos, and structured data files.
"""

import os
import json
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Dict, List, Optional, Union, Any
import base64
import mimetypes

# Document processing
try:
    import docx
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

# Image processing
try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# OpenAI for vision analysis
try:
    from openai import OpenAI
    OPENAI_VISION_AVAILABLE = True
except ImportError:
    OPENAI_VISION_AVAILABLE = False


class FileAnalysisResult:
    """Container for file analysis results."""
    
    def __init__(self, file_path: str, file_type: str, content: str, metadata: Dict[str, Any]):
        self.file_path = file_path
        self.file_type = file_type
        self.content = content
        self.metadata = metadata
        self.summary = self._generate_summary()
    
    def _generate_summary(self) -> str:
        """Generate a summary of the file content."""
        if not self.content:
            return "No content extracted from file."
        
        # Truncate content for summary
        content_preview = self.content[:500] + "..." if len(self.content) > 500 else self.content
        return f"File type: {self.file_type}\nContent preview: {content_preview}"
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for serialization."""
        return {
            "file_path": self.file_path,
            "file_type": self.file_type,
            "content": self.content,
            "metadata": self.metadata,
            "summary": self.summary
        }


class FileAnalyzer:
    """Comprehensive file analyzer for various file types."""
    
    def __init__(self, openai_api_key: Optional[str] = None):
        """
        Initialize the file analyzer.
        
        Args:
            openai_api_key: OpenAI API key for vision analysis (optional)
        """
        self.openai_client = None
        if openai_api_key and OPENAI_VISION_AVAILABLE:
            self.openai_client = OpenAI(api_key=openai_api_key)
    
    def analyze_file(self, file_path: Union[str, Path]) -> FileAnalysisResult:
        """
        Analyze a file and extract relevant information.
        
        Args:
            file_path: Path to the file to analyze
            
        Returns:
            FileAnalysisResult containing extracted content and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file type
        mime_type, _ = mimetypes.guess_type(str(file_path))
        file_extension = file_path.suffix.lower()
        
        # Analyze based on file type
        if file_extension in ['.txt', '.md', '.markdown']:
            return self._analyze_text_file(file_path, mime_type)
        elif file_extension in ['.pdf']:
            return self._analyze_pdf_file(file_path, mime_type)
        elif file_extension in ['.docx', '.doc']:
            return self._analyze_docx_file(file_path, mime_type)
        elif file_extension in ['.xlsx', '.xls', '.csv']:
            return self._analyze_spreadsheet_file(file_path, mime_type)
        elif file_extension in ['.json']:
            return self._analyze_json_file(file_path, mime_type)
        elif file_extension in ['.xml']:
            return self._analyze_xml_file(file_path, mime_type)
        elif file_extension in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']:
            return self._analyze_image_file(file_path, mime_type)
        elif file_extension in ['.mp4', '.avi', '.mov', '.wmv']:
            return self._analyze_video_file(file_path, mime_type)
        else:
            return self._analyze_unknown_file(file_path, mime_type)
    
    def _analyze_text_file(self, file_path: Path, mime_type: str) -> FileAnalysisResult:
        """Analyze text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            metadata = {
                "file_size": file_path.stat().st_size,
                "encoding": "utf-8",
                "line_count": len(content.splitlines()),
                "word_count": len(content.split())
            }
            
            return FileAnalysisResult(
                file_path=str(file_path),
                file_type="text",
                content=content,
                metadata=metadata
            )
        except Exception as e:
            return FileAnalysisResult(
                file_path=str(file_path),
                file_type="text",
                content=f"Error reading text file: {str(e)}",
                metadata={"error": str(e)}
            )
    
    def _analyze_pdf_file(self, file_path: Path, mime_type: str) -> FileAnalysisResult:
        """Analyze PDF files."""
        content = ""
        metadata = {"file_size": file_path.stat().st_size}
        
        # Try pdfplumber first (better for complex layouts)
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(file_path) as pdf:
                    content = "\n".join([page.extract_text() or "" for page in pdf.pages])
                    metadata["page_count"] = len(pdf.pages)
            except Exception as e:
                metadata["pdfplumber_error"] = str(e)
        
        # Fallback to PyPDF2
        if not content and PDF_AVAILABLE:
            try:
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    content = "\n".join([page.extract_text() for page in pdf_reader.pages])
                    metadata["page_count"] = len(pdf_reader.pages)
            except Exception as e:
                metadata["pypdf2_error"] = str(e)
        
        if not content:
            content = "Could not extract text from PDF file."
        
        return FileAnalysisResult(
            file_path=str(file_path),
            file_type="pdf",
            content=content,
            metadata=metadata
        )
    
    def _analyze_docx_file(self, file_path: Path, mime_type: str) -> FileAnalysisResult:
        """Analyze Word documents."""
        if not DOCX_AVAILABLE:
            return FileAnalysisResult(
                file_path=str(file_path),
                file_type="docx",
                content="python-docx library not available. Install with: pip install python-docx",
                metadata={"error": "python-docx not available"}
            )
        
        try:
            doc = Document(file_path)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Extract table content
            table_content = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    table_content.append(row_text)
            
            if table_content:
                content += "\n\nTables:\n" + "\n".join(table_content)
            
            metadata = {
                "file_size": file_path.stat().st_size,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
            
            return FileAnalysisResult(
                file_path=str(file_path),
                file_type="docx",
                content=content,
                metadata=metadata
            )
        except Exception as e:
            return FileAnalysisResult(
                file_path=str(file_path),
                file_type="docx",
                content=f"Error reading DOCX file: {str(e)}",
                metadata={"error": str(e)}
            )
    
    def _analyze_spreadsheet_file(self, file_path: Path, mime_type: str) -> FileAnalysisResult:
        """Analyze Excel and CSV files."""
        if not PANDAS_AVAILABLE:
            return FileAnalysisResult(
                file_path=str(file_path),
                file_type="spreadsheet",
                content="pandas library not available. Install with: pip install pandas openpyxl",
                metadata={"error": "pandas not available"}
            )
        
        try:
            if file_path.suffix.lower() == '.csv':
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            
            # Get basic info
            content = f"Spreadsheet with {len(df)} rows and {len(df.columns)} columns\n\n"
            content += f"Columns: {', '.join(df.columns.tolist())}\n\n"
            
            # Show first few rows
            content += "Sample data (first 5 rows):\n"
            content += df.head().to_string()
            
            # Show data types
            content += f"\n\nData types:\n{df.dtypes.to_string()}"
            
            # Show summary statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                content += f"\n\nSummary statistics:\n{df[numeric_cols].describe().to_string()}"
            
            metadata = {
                "file_size": file_path.stat().st_size,
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": df.columns.tolist(),
                "data_types": df.dtypes.to_dict()
            }
            
            return FileAnalysisResult(
                file_path=str(file_path),
                file_type="spreadsheet",
                content=content,
                metadata=metadata
            )
        except Exception as e:
            return FileAnalysisResult(
                file_path=str(file_path),
                file_type="spreadsheet",
                content=f"Error reading spreadsheet file: {str(e)}",
                metadata={"error": str(e)}
            )
    
    def _analyze_json_file(self, file_path: Path, mime_type: str) -> FileAnalysisResult:
        """Analyze JSON files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Pretty print the JSON
            content = json.dumps(data, indent=2, ensure_ascii=False)
            
            # Extract structure info
            structure_info = self._analyze_json_structure(data)
            
            metadata = {
                "file_size": file_path.stat().st_size,
                "structure": structure_info
            }
            
            return FileAnalysisResult(
                file_path=str(file_path),
                file_type="json",
                content=content,
                metadata=metadata
            )
        except Exception as e:
            return FileAnalysisResult(
                file_path=str(file_path),
                file_type="json",
                content=f"Error reading JSON file: {str(e)}",
                metadata={"error": str(e)}
            )
    
    def _analyze_json_structure(self, data: Any, path: str = "") -> Dict[str, Any]:
        """Analyze JSON structure recursively."""
        if isinstance(data, dict):
            return {
                "type": "object",
                "keys": list(data.keys()),
                "properties": {k: self._analyze_json_structure(v, f"{path}.{k}") for k, v in data.items()}
            }
        elif isinstance(data, list):
            if data:
                return {
                    "type": "array",
                    "length": len(data),
                    "item_type": type(data[0]).__name__,
                    "sample_item": self._analyze_json_structure(data[0], f"{path}[0]")
                }
            else:
                return {"type": "array", "length": 0}
        else:
            return {"type": type(data).__name__, "value": str(data)[:100]}
    
    def _analyze_xml_file(self, file_path: Path, mime_type: str) -> FileAnalysisResult:
        """Analyze XML files."""
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            # Extract text content
            content = ET.tostring(root, encoding='unicode', method='text')
            
            # Analyze structure
            structure = self._analyze_xml_structure(root)
            
            metadata = {
                "file_size": file_path.stat().st_size,
                "root_tag": root.tag,
                "structure": structure
            }
            
            return FileAnalysisResult(
                file_path=str(file_path),
                file_type="xml",
                content=content,
                metadata=metadata
            )
        except Exception as e:
            return FileAnalysisResult(
                file_path=str(file_path),
                file_type="xml",
                content=f"Error reading XML file: {str(e)}",
                metadata={"error": str(e)}
            )
    
    def _analyze_xml_structure(self, element: ET.Element) -> Dict[str, Any]:
        """Analyze XML structure recursively."""
        children = list(element)
        return {
            "tag": element.tag,
            "attributes": element.attrib,
            "text": element.text.strip() if element.text else None,
            "children": [self._analyze_xml_structure(child) for child in children]
        }
    
    def _analyze_image_file(self, file_path: Path, mime_type: str) -> FileAnalysisResult:
        """Analyze image files using OCR and vision analysis."""
        content = ""
        metadata = {"file_size": file_path.stat().st_size}
        
        # Get image dimensions
        try:
            with Image.open(file_path) as img:
                metadata["dimensions"] = img.size
                metadata["format"] = img.format
                metadata["mode"] = img.mode
        except Exception as e:
            metadata["image_error"] = str(e)
        
        # OCR analysis
        if OCR_AVAILABLE:
            try:
                with Image.open(file_path) as img:
                    ocr_text = pytesseract.image_to_string(img)
                    if ocr_text.strip():
                        content += f"OCR Text:\n{ocr_text}\n\n"
            except Exception as e:
                metadata["ocr_error"] = str(e)
        
        # Vision analysis with OpenAI
        if self.openai_client:
            try:
                with open(file_path, "rb") as image_file:
                    base64_image = base64.b64encode(image_file.read()).decode('utf-8')
                
                response = self.openai_client.chat.completions.create(
                    model="gpt-4-vision-preview",
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": "Analyze this image and describe what you see. If it's a UI mockup, wireframe, or diagram, provide detailed information about the interface elements, layout, and functionality."
                                },
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/jpeg;base64,{base64_image}"
                                    }
                                }
                            ]
                        }
                    ],
                    max_tokens=1000
                )
                
                vision_analysis = response.choices[0].message.content
                content += f"Vision Analysis:\n{vision_analysis}\n"
                
            except Exception as e:
                metadata["vision_error"] = str(e)
        
        if not content:
            content = "Image file - no text content extracted."
        
        return FileAnalysisResult(
            file_path=str(file_path),
            file_type="image",
            content=content,
            metadata=metadata
        )
    
    def _analyze_video_file(self, file_path: Path, mime_type: str) -> FileAnalysisResult:
        """Analyze video files."""
        # For now, we'll just provide basic file information
        # In a full implementation, you might want to extract frames and analyze them
        metadata = {
            "file_size": file_path.stat().st_size,
            "note": "Video analysis not fully implemented. Consider extracting key frames for analysis."
        }
        
        content = f"Video file: {file_path.name}\n"
        content += f"Size: {file_path.stat().st_size / (1024*1024):.2f} MB\n"
        content += "Note: Full video analysis requires additional processing to extract frames."
        
        return FileAnalysisResult(
            file_path=str(file_path),
            file_type="video",
            content=content,
            metadata=metadata
        )
    
    def _analyze_unknown_file(self, file_path: Path, mime_type: str) -> FileAnalysisResult:
        """Handle unknown file types."""
        metadata = {
            "file_size": file_path.stat().st_size,
            "mime_type": mime_type,
            "note": "Unknown file type - limited analysis available"
        }
        
        content = f"Unknown file type: {file_path.suffix}\n"
        content += f"MIME type: {mime_type}\n"
        content += f"Size: {file_path.stat().st_size} bytes"
        
        return FileAnalysisResult(
            file_path=str(file_path),
            file_type="unknown",
            content=content,
            metadata=metadata
        )
    
    def analyze_multiple_files(self, file_paths: List[Union[str, Path]]) -> List[FileAnalysisResult]:
        """
        Analyze multiple files and return results.
        
        Args:
            file_paths: List of file paths to analyze
            
        Returns:
            List of FileAnalysisResult objects
        """
        results = []
        for file_path in file_paths:
            try:
                result = self.analyze_file(file_path)
                results.append(result)
            except Exception as e:
                # Create error result
                error_result = FileAnalysisResult(
                    file_path=str(file_path),
                    file_type="error",
                    content=f"Error analyzing file: {str(e)}",
                    metadata={"error": str(e)}
                )
                results.append(error_result)
        
        return results
    
    def generate_combined_analysis(self, results: List[FileAnalysisResult]) -> str:
        """
        Generate a combined analysis summary from multiple file results.
        
        Args:
            results: List of FileAnalysisResult objects
            
        Returns:
            Combined analysis string
        """
        if not results:
            return "No files analyzed."
        
        combined = "=== FILE ANALYSIS SUMMARY ===\n\n"
        
        for i, result in enumerate(results, 1):
            combined += f"File {i}: {Path(result.file_path).name}\n"
            combined += f"Type: {result.file_type}\n"
            combined += f"Summary: {result.summary}\n"
            combined += "-" * 50 + "\n\n"
        
        # Extract key insights
        combined += "=== KEY INSIGHTS ===\n\n"
        
        # Group by file type
        by_type = {}
        for result in results:
            if result.file_type not in by_type:
                by_type[result.file_type] = []
            by_type[result.file_type].append(result)
        
        for file_type, type_results in by_type.items():
            combined += f"{file_type.upper()} FILES ({len(type_results)}):\n"
            for result in type_results:
                combined += f"  - {Path(result.file_path).name}: {result.content[:100]}...\n"
            combined += "\n"
        
        return combined
